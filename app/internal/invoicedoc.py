from docx import Document
from typing import List
from docx.shared import Pt

from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_invoice(invoice, customer, orders: List, products: dict):
    doc = Document()
    title = doc.add_heading(f"{customer.store_name}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading("Invoice", 0)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Customer: {customer.company_name}")
    doc.add_paragraph(f"Email: {customer.email}")
    doc.add_paragraph(f"Date: {invoice.date.strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"Status: {invoice.status}")
    doc.add_paragraph(" ")

    doc.add_heading("Products", level=1)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Product Name"
    hdr_cells[1].text = "Unit Price"
    hdr_cells[2].text = "Quantity"
    hdr_cells[3].text = "Tax (%)"
    hdr_cells[4].text = "Total"

    total_amount = 0

    for order in orders:
        product = products.get(order.product_id)
        if not product:
            continue

        unit_price = product.unit_price
        quantity = order.quantity
        tax_percent = product.tax_percent or 0

        subtotal = unit_price * quantity
        total_amount += subtotal

        row_cells = table.add_row().cells
        row_cells[0].text = product.name
        row_cells[1].text = f"${unit_price:.2f}"
        row_cells[2].text = str(quantity)
        row_cells[3].text = f"{tax_percent:.2f}%"
        row_cells[4].text = f"${subtotal:.2f}"

    empty_row = table.add_row().cells
    first_cell = empty_row[0]
    for cell in empty_row[1:]:
        first_cell.merge(cell)
    first_cell.text = ""

    def add_summary_row(table, label, value):
        row = table.add_row().cells
        left_cell = row[0]
        right_cell = row[-1]
        if len(row) > 2:
            left_cell.merge(row[1])
            for i in range(2, len(row) - 1):
                left_cell.merge(row[i])

        left_cell.text = label
        left_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        right_cell.text = value
        for paragraph in right_cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                font = run.font
                font.size = Pt(11)

    add_summary_row(table, "Subtotal", f"${total_amount:.2f}")
    add_summary_row(table, "Discount", f"{invoice.discount:.2f}%")
    total_after_discount = total_amount - (total_amount * invoice.discount / 100)
    add_summary_row(table, "Total", f"${total_after_discount:.2f}")
    
    return doc